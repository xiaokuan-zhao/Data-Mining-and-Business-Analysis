import os
import shutil
from datetime import datetime

from docx import Document
import numpy as np
import pandas as pd
import re
import time
from tqdm import tqdm

# 读入 需要处理文件名 的xlsx文件
df = pd.read_excel('C:/Users/zhong/Desktop/111.xlsx', sheet_name='Sheet1', usecols=[0])
data = df.values.tolist()
correct_files_name = [n[0] + ".docx" for n in data]  # 加入".docx"后缀

path = "C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word/"  # 文件夹目录
correct_results = []
for n in tqdm(correct_files_name):
    document = Document(path + n)
    s = ""
    for paragraph in document.paragraphs:
        s = s + paragraph.text
    s = s.replace(' ', '')
    s = s.replace(",", "，")
    s = s.replace("(", "（")
    s = s.replace(")", "）")
    s = s.replace(":", "：")
    s = s.replace("\u3000", "")
    s = s.replace("\xa0", "")
    try:
        cssj = re.findall("，([0-9]{4})年", s)
        cssj = [int(n) for n in cssj]
        correct_results.append(min(cssj))
    except:
        correct_results.append("")
    time.sleep(0.05)

# 根据只言片语判断精确地址
import cpca

# 读入地址
df = pd.read_excel('C:/Users/zhong/Desktop/111.xlsx', sheet_name='Sheet1', usecols=[0])
data = df.values.tolist()
data = [n[0] for n in data]
locates = cpca.transform(data)
df_locates = pd.DataFrame(locates)
path1 = "C:/Users/zhong/Desktop/"  # 保存至csv
df_locates.to_csv(path1 + '数据修正结果2.csv', encoding="gbk")

# 预测法官性别
import ngender

df = pd.read_excel('C:/Users/zhong/Desktop/111.xlsx', sheet_name='Sheet1', usecols=[0])
data = df.values.tolist()
data = [n[0] for n in data]
gender = []
for n in data:
    gender.append(ngender.guess(n))

# 将中文日期转换成标准模式
df = pd.read_excel(r"C:/Users/zhong/Desktop/111.xlsx")


def func(x):
    try:
        year = x.split("年")[0]
        month = x.split("年")[1].split("月")[0]
        day = x.split("年")[1].split("月")[1].split("日")[0]
        if len(day) >= 3:
            day = day[0] + day[2]
        chinese_english = dict(〇=0, 一=1, 二=2, 三=3, 四=4, 五=5, 六=6, 七=7, 八=8, 九=9, 十=10)
        year = "".join(str(chinese_english[i]) for i in year)
        month = "".join(str(chinese_english[i]) for i in month)
        day = "".join(str(chinese_english[i]) for i in day)
        if len(month) == 3:
            month = month[0] + month[2]
        if len(day) == 3:
            day = day[0] + day[2]
        final_date = year + "." + month + "." + day
        return final_date
    except:
        return ""


df["final_date"] = df["判决时间"].apply(func)
path1 = "C:/Users/zhong/Desktop/"  # 保存至csv
df.to_csv(path1 + '数据修正结果1.csv', encoding="gbk")

# 判处结果规范化
df = pd.read_excel('C:/Users/zhong/Desktop/111.xlsx', sheet_name='Sheet1', usecols=[0])
data = df.values.tolist()
data = [n[0] for n in data]
pj = []
for s in data:
    d = {}
    try:
        jy = re.findall("(拘役[一二两三四五六七八九十个半月年日天零又]*)", s)[0]
        d["拘役"] = jy
    except:
        d["拘役"] = ""
    try:
        yqtx = re.findall("(有期徒刑[一二两三四五六七八九十个半月年日天零又]*?)[，。缓;；、并]", s)[0]
        d["有期徒刑"] = yqtx
    except:
        d["有期徒刑"] = ""
    try:
        hx = re.findall("(缓刑[一二两三四五六七八九十个半月年日天零又]*)", s)[0]
        d["缓刑"] = hx
    except:
        d["缓刑"] = ""
    try:
        fj = re.findall("([一二两三四五六七八九十百千万0-9oOl1]*?元)", s)[0]
        d["罚金"] = fj
    except:
        d["罚金"] = ""
    pj.append(d)

df_files = pd.DataFrame(pj)
path1 = "C:/Users/zhong/Desktop/"  # 保存至csv
df_files.to_csv(path1 + '数据修正结果2.csv', encoding="gbk")

# 判决结果再修正
df = pd.read_excel('C:/Users/zhong/Desktop/111.xlsx', sheet_name='Sheet1', usecols=[0])
data = df.values.tolist()
data = [n[0] for n in data]
d = []
for n in data:
    s = []
    try:
        s = n.split("m")
    except:
        s = ["", ""]
        print(n)
    d.append(s)
df_files = pd.DataFrame(d)
path1 = "C:/Users/zhong/Desktop/"  # 保存至csv
df_files.to_csv(path1 + '数据修正结果2.csv', encoding="gbk")

# 罚金修正
df = pd.read_excel('C:/Users/zhong/Desktop/111.xlsx', sheet_name='Sheet1', usecols=[0])
data = df.values.tolist()
data = [n[0] for n in data]
d = []
for n in data:
    fj = re.findall("([一二两三四五六七八九十百千万0-9oOl1,.，３壹贰叁肆伍陆柒捌玖拾佰仟２]*?元)", n)[0]
    d.append(fj)
df_files = pd.DataFrame(d)
path1 = "C:/Users/zhong/Desktop/"  # 保存至csv
df_files.to_csv(path1 + '数据修正结果2.csv', encoding="gbk")

# 算年龄（有点问题）
from datetime import datetime

df = pd.read_excel('C:/Users/zhong/Desktop/111.xlsx', sheet_name='Sheet1', usecols=[0, 1])
def calculate_age(pjsj, birth):
    try:
        birth_d = datetime.datetime.strptime(birth, "%Y-%m-%d")
        today_d = datetime.datetime.strptime(pjsj, "%Y-%m-%d")
        birth_t = birth_d.replace(year=today_d.year)
        if today_d > birth_t:
            age = today_d.year - birth_d.year
        else:
            age = today_d.year - birth_d.year - 1
        return age
    except:
        return ""

df["时间差"] = df["判决时间", "出生时间"].apply(calculate_age)
path1 = "C:/Users/zhong/Desktop/"  # 保存至csv
df.to_csv(path1 + '数据修正结果2.csv', encoding="gbk")

# 缓刑修正
df = pd.read_excel('C:/Users/zhong/Desktop/perfect3.xlsx', sheet_name='Sheet1', usecols=[1])
data = df.values.tolist()
correct_files_name = [n[0] + ".docx" for n in data]  # 加入".docx"后缀

path = "C:/Users/zhong/Desktop/危险驾驶罪项目/数据挖掘与商务分析数据/total/"  # 文件夹目录
correct_results = []
for n in tqdm(correct_files_name):
    document = Document(path + n)
    s = ""
    for paragraph in document.paragraphs:
        s = s + paragraph.text
    try:
        s = s.replace(' ', '')
        s = s.replace(",", "，")
        s = s.replace("(", "（")
        s = s.replace(")", "）")
        s = s.replace(":", "：")
        s = s.replace("\u3000", "")
        s = s.replace("\xa0", "")
        s = s.split("二〇")[0]  # 把附属法条的部分都删去
    except:
        print(n+"***1")
    try:
        s = s.split("判决如下")[1]
    except:
        print(n+"***2")
    try:
        ifhx = re.findall("缓.*", s)[0]
        correct_results.append([ifhx, 1])
    except:
        correct_results.append(["", 0])
    time.sleep(0.05)

df_files = pd.DataFrame(correct_results)
path1 = "C:/Users/zhong/Desktop/"  # 保存至csv
df_files.to_csv(path1 + '数据修正结果3.csv', encoding="gbk")

correct_results[756]