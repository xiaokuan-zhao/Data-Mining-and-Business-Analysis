import os
import shutil
from docx import Document
import numpy as np
import pandas as pd
import re
import time
from tqdm import tqdm

path = "C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word/"  # 文件夹目录
path2 = "C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word_1/"  # 结果目录
files = os.listdir(path)  # 得到文件夹下的所有文件名称
files1 = list(filter(lambda f: f.find("危险驾驶罪") != -1, files))  # 提取出危险驾驶罪的word文档
print(files1[:10])
wordfile_name = list(filter(lambda f: f.find("某") == -1, files1))  # 把"某某"的人给去除
print(wordfile_name[:10])

shutil.move('C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word/1002余建辉危险驾驶罪一审刑事判决书.docx', path2)
for n in wordfile_name:
    shutil.copyfile(path + n, path2 + n)

# 创建个测试集来试一试
path3 = "C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word_test/"
wordfile_name_test = wordfile_name[::200]
for n in wordfile_name_test:
    shutil.copyfile(path2 + n, path3 + n)


# 定义一个函数:处理已经提取出来的字符串形式的判决书文本

def text_mining(n, s):
    # 去除或规范特殊符号
    s = s.replace(' ', '')
    s = s.replace(",", "，")
    s = s.replace("(", "（")
    s = s.replace(")", "）")
    s = s.replace(":", "：")
    s = s.replace("\u3000", "")
    s = s.replace("\xa0", "")
    d = {}  # 提取数据特征构成字典
    # ### 法院信息 ###
    # 案件名称
    ajmc = n.split(".")[0]  # 去除.docx
    d["案件名称"] = ajmc
    # 几审
    js = re.findall(".审", ajmc)[0]
    d["几审"] = js
    # 法院所在地
    fyszd = re.findall("(.*?)人民法院", s)[0]
    d["法院所在地"] = fyszd
    # 审判员姓名
    fgxm = re.findall("审判[长员](.*?)[二审人]", s)
    d["法官姓名"] = min(fgxm, key=len, default='')
    # 判决时间
    time = re.findall("二〇.*?日", s)[0]
    d["判决时间"] = time
    # ### 被告人基本信息 ###
    # 民族
    mz = re.findall("[\u4e00-\u9fa5]*?族", s)
    mz = min(mz, key=len, default='')
    d["民族"] = mz
    # 姓名
    xm = re.findall("被告人(.*?)，", s)[0]
    d["姓名"] = xm
    # 性别
    xb = re.findall("，([男女])，", s)[0]
    d["性别"] = xb
    # 学历
    xl = re.findall("，(文盲|[\u4e00-\u9fa5]*?文化)，", s)[0]
    d["学历"] = xl
    # 出生地/户籍地
    cs = re.findall("出生于([\u4e00-\u9fa5]*?)，|户籍地([\u4e00-\u9fa5]*?)，|户籍所在地([\u4e00-\u9fa5]*?)，"
                    "|籍贯([\u4e00-\u9fa5]*?)，", s)[0]
    hj = max(cs, key=len, default='')
    d["户籍"] = hj
    # 出生时间
    cssj = re.findall("，([0-9]{4}年.*?日.*?)", s)[0]
    d["出生时间"] = cssj
    # ### 案件经过 ###
    # 血液乙醇浓度
    ycnd = re.findall("[0-9.oOl±()]*?[mｍ][gｇ][/／∕][0-9.]*?[mｍ][lｌL1Ｌ]|[0-9.±oOl()]*?毫克[/／∕][0-9.]*?毫升", s)
    ycnd = [x for x in ycnd if x != "80mg/100ml" and x != "80毫克/100毫升" and x != "200毫克/100毫升" and x != "200mg/100ml"
            and x != "200mg" and x != "2mg" and x != "200毫克" and x != "2毫克" and x != ""]
    d["血液乙醇浓度"] = ycnd[-1]
    # 是否如实供述/坦白
    s1 = s.split("二〇")[0]  # 把附属法条的部分都删去
    rsgs = 0
    if "如实供述" in s1 or "坦白" in s1:
        rsgs = 1
    d["如实供述"] = rsgs
    # 是否认罪认罚
    rzrf = 0
    if "认罪认罚" in s1:
        rzrf = 1
    d["认罪认罚"] = rzrf
    # 是否有人顶替
    dt = 0
    if "顶替" in s1:
        dt = 1
    d["顶替"] = dt
    # 是否逃逸
    ty = 0
    if "逃逸" in s1:
        ty = 1
    d["逃逸"] = ty
    # 是否在高速道路
    gsdl = 0
    if "高速公路" in s1 or "城市快速路" in s1:
        gsdl = 1
    d["高速道路"] = gsdl
    # 是否与事故有关
    sg = 0
    if "事故" in s1:
        sg = 1
    d["事故"] = sg
    # 对方伤亡情况(无伤害取0，轻伤取1，重伤取2，死亡取3)
    swqk = 0
    if "轻伤" in s1:
        swqk = 1
        if "重亡" in s1:
            swqk = 2
            if "死亡" in s1:
                swqk = 3
    d["伤亡情况"] = swqk
    # 事故是否全责
    qbzr = 0
    if "全部责任" in s1 or "全责" in s1:
        qbzr = 1
    d["全部责任"] = qbzr
    # 是否自首
    zs = 0
    if "自首" in s1:
        zs = 1
    d["自首"] = zs
    # 是否谅解
    lj = 0
    if "谅解" in s1:
        lj = 1
    d["谅解"] = lj
    # 是否无证驾驶
    wzjs = 0
    if "无证驾驶" in s1 or "无资格" in s1 or "无驾驶" in s1 or "取得驾驶资格" in s1:
        wzjs = 1
    d["无证驾驶"] = wzjs
    # 是否有酒驾犯罪史
    jjfzs = 0
    if "曾因酒" in s1:
        jjfzs = 1
    d["酒驾犯罪史"] = jjfzs

    # ### 判决结果 ###
    # 判决结果
    pjjg = re.findall("判决如下：(.*?)（", s)[0]
    d["判决结果"] = pjjg

    return d


# 开始遍历测试集，运用定义函数清洗数据
wrong_file_name = []  # 乱码文件名
complex_file_name = []  # 需要人工提取的文件名
data_mining_result = []  # 提取的数据特征字典
for n in wordfile_name_test:
    try:
        document = Document(path3 + n)
        s = ""
        for paragraph in document.paragraphs:
            s = s + paragraph.text
        if s != "ȏׯȏɽЌ ʂ Ő ":
            d = text_mining(n, s)
            data_mining_result.append(d)
        else:
            wrong_file_name.append(n)
    except:
        complex_file_name.append(n)

print(wrong_file_name)
print(complex_file_name)

for d in data_mining_result:
    print(d)

# 开始遍历全部集，运用定义函数清洗数据
wrong_file_name = []  # 乱码文件名
complex_file_name = []  # 需要人工提取的文件名
data_mining_result = []  # 提取的数据特征字典
for n in tqdm(files1):
    try:
        document = Document(path2 + n)
        s = ""
        for paragraph in document.paragraphs:
            s = s + paragraph.text
        if s != "ȏׯȏɽЌ ʂ Ő ":
            d = text_mining(n, s)
            data_mining_result.append(d)
        else:
            wrong_file_name.append(n)
    except:
        complex_file_name.append(n)
    time.sleep(0.05)

# print(wrong_file_name)
# print(complex_file_name)
# print(data_mining_result)

print(len(wrong_file_name))
print(len(complex_file_name))
print(len(data_mining_result))

# 转移到各自的文件夹
path = "C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word/"  # 文件夹目录
path2 = "C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word_1/"  # 危险驾驶罪所在原始目录(去除某某)
pathw = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/wrong/"  # 错误文件的转移目录
pathp = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/perfect/"  # 完美文件的转移目录
pathc = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/complex/"  # 复杂文件的转移目录
pathx = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/mou/"  # 某某文件的转移目录
for n in wrong_file_name:
    shutil.copyfile(path2 + n, pathw + n)
c = []  # c包含某某的所有文件
for n in complex_file_name:
    try:
        shutil.copyfile(path2 + n, pathc + n)
    except:
        c.append(n)
for n in files1:
    if n not in wrong_file_name and n not in complex_file_name:
        shutil.copyfile(path2 + n, pathp + n)
for n in files1:
    if n in c:
        shutil.copyfile(path + n, pathx + n)

# perfect文件转换成dataframe形式
df_perfectfiles = pd.DataFrame(data_mining_result)
path4 = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/"  # 保存至csv
df_perfectfiles.to_csv(path4+'perfectfiles.csv',encoding="gbk")

# ###### 再次操作程序
import os
import shutil
from docx import Document
import numpy as np
import pandas as pd
import re
import time
from tqdm import tqdm

path = "C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word/"  # 文件夹目录
path2 = "C:/Users/zhong/Desktop/危险驾驶罪项目/原始数据/word_1/"  # 危险驾驶罪所在原始目录(去除某某)
pathw = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/wrong/"  # 错误文件的转移目录
pathp = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/perfect/"  # 完美文件的转移目录
pathc = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/complex/"  # 复杂文件的转移目录
pathx = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/mou/"  # 某某文件的转移目录
pathxc = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/mou_c/"
pathxp = "C:/Users/zhong/Desktop/危险驾驶罪项目/中间过程/mou_p/"
files = os.listdir(pathw)
wrong_file_name = list(files)
files = os.listdir(pathp)
perfect_file_name = list(files)
files = os.listdir(pathc)
complex_file_name = list(files)
files = os.listdir(pathxc)
mou_complex_files_name = list(files)
files = os.listdir(pathxp)
mou_perfect_files_name = list(files)

w = []
c = []
data_mining_result = []  # 提取的数据特征字典
for n in tqdm(mou_complex_files_name):
    try:
        document = Document(pathxc + n)
        s = ""
        for paragraph in document.paragraphs:
            s = s + paragraph.text
        if s != "ȏׯȏɽЌ ʂ Ő ":
            d = text_mining(n, s)
            data_mining_result.append(d)
        else:
            w.append(n)
    except:
        c.append(n)
    time.sleep(0.05)


for n in mou_complex_files_name[:100]:
    document = Document(pathxc + n)
    s = ""
    for paragraph in document.paragraphs:
        s = s + paragraph.text
    try:
        print(text_mining(n, s))
    except:
        print(s)

